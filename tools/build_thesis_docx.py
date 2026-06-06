import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "HOLMES_IDS_THESIS_WORKING_DRAFT.md"
OUTPUT = ROOT / "HOLMES_IDS_THESIS_WORKING_DRAFT.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x20, 0x20, 0x20)
LIGHT_FILL = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(i, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean_inline(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def clean_inline(text):
    text = text.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("Heading 4", 11, DARK_BLUE, 6, 3),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("HOLMES IDS\nThesis Working Draft")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    r = subtitle.add_run("Hybrid Online Learning Model for Enhanced Security")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.color.rgb = DARK_BLUE

    meta = [
        ("Document type", "Final Year Graduation Project Thesis - Working Draft"),
        ("Authors", "Mohamed Ahmed Abdelfattah; Mazen Ibrahim Abdelrazek; Mohamed Abdelgawad Abdelrahman; Hala Mazen Waddad; Sohaila Mustafa Abdelfattah"),
        ("Supervisor", "Dr. Soha Ehsan"),
        ("Institution", "[INSTITUTION NAME - VERIFY]"),
        ("Department", "[DEPARTMENT NAME - VERIFY]"),
        ("Planning date", "2026-06-06"),
        ("Source file", "HOLMES_IDS_THESIS_WORKING_DRAFT.md"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 7160])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_shading(row.cells[0], LIGHT_FILL)
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Note: Items marked VERIFY, INSERT, or TO BE COMPLETED must be confirmed before final submission.").italic = True
    doc.add_page_break()


def add_header_footer(doc):
    for section in doc.sections:
        header_p = section.header.paragraphs[0]
        header_p.text = "HOLMES IDS Thesis Working Draft"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header_p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.text = "Working draft for review and formatting"
        for run in footer_p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def is_table_line(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def column_widths(cols):
    total = 9360
    if cols <= 0:
        return []
    if cols == 2:
        return [2600, total - 2600]
    if cols == 3:
        return [1800, 3600, total - 5400]
    if cols == 4:
        return [1300, 2500, 2800, total - 6600]
    width = total // cols
    widths = [width] * cols
    widths[-1] += total - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    parsed = [split_table_row(r) for r in rows if not re.match(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", r.strip())]
    if not parsed:
        return
    max_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    table.style = "Table Grid"
    set_table_geometry(table, column_widths(max_cols))
    for r_idx, row_data in enumerate(parsed):
        for c_idx in range(max_cols):
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = table.rows[r_idx].cells[c_idx]
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            set_cell_text(cell, text, bold=(r_idx == 0))
    doc.add_paragraph()


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code.rstrip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(doc, text, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(clean_inline(text))


def add_paragraph_with_inline(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    text = clean_inline(text)
    if text.startswith("`[") and text.endswith("]`"):
        run = p.add_run(text.strip("`"))
        run.italic = True
        run.font.color.rgb = RGBColor(0x9B, 0x1C, 0x1C)
    else:
        p.add_run(text)


def build_docx():
    source_text = SOURCE.read_text(encoding="utf-8")
    lines = source_text.splitlines()

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    in_code = False
    code_lines = []
    table_lines = []
    skipped_first_h1 = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                if table_lines:
                    add_markdown_table(doc, table_lines)
                    table_lines = []
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if is_table_line(line):
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            if not skipped_first_h1:
                skipped_first_h1 = True
            else:
                doc.add_heading(clean_inline(stripped[2:]), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(clean_inline(stripped[5:]), level=3)
        elif re.match(r"^\d+\.\s+", stripped):
            add_bullet(doc, re.sub(r"^\d+\.\s+", "", stripped), ordered=True)
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:], ordered=False)
        else:
            add_paragraph_with_inline(doc, stripped)
        i += 1

    if table_lines:
        add_markdown_table(doc, table_lines)
    if code_lines:
        add_code_block(doc, "\n".join(code_lines))

    add_header_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
