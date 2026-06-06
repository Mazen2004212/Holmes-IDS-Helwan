from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "HOLMES_IDS_Final_Documentation_REFERENCE_TEMPLATE.docx"
SCREENSHOTS = ROOT / "presentation_assets" / "screenshots"


TEAM = [
    "Mohamed Ahmed Abdelfattah",
    "Mazen Ibrahim Abdelrazek",
    "Mohamed Abdelgawad Abdelrahman",
    "Hala Mazen Waddad",
    "Sohaila Mustafa Abdelfattah",
]


page_counter = 1


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(4)


def run(p, text, size=12, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return r


def p(doc, text="", size=12, bold=False, italic=False, align=None, before=0, after=4):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(before)
    par.paragraph_format.space_after = Pt(after)
    par.paragraph_format.line_spacing = 1.15
    if align is not None:
        par.alignment = align
    run(par, text, size=size, bold=bold, italic=italic)
    return par


def bullet(doc, text):
    par = p(doc, "", after=3)
    par.paragraph_format.left_indent = Inches(0.25)
    run(par, "➢ ", size=12)
    run(par, text, size=12)


def dot_bullet(doc, text):
    par = p(doc, "", after=2)
    par.paragraph_format.left_indent = Inches(0.35)
    run(par, "• ", size=12)
    run(par, text, size=12)


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def page_no(doc):
    global page_counter
    par = p(doc, f"[{page_counter}]", align=WD_ALIGN_PARAGRAPH.LEFT, after=12)
    page_counter += 1
    return par


def new_page(doc):
    doc.add_page_break()


def blank(doc):
    new_page(doc)
    new_page(doc)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(cell.paragraphs[0], text, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run(cells[i].paragraphs[0], text)
    p(doc, "", after=8)
    return t


def figure_placeholder(doc, caption, image_name=None):
    for _ in range(7):
        p(doc, "")
    if image_name:
        path = SCREENSHOTS / image_name
        if path.exists():
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.add_run().add_picture(str(path), width=Inches(5.8))
    else:
        p(doc, "[Insert Diagram Here]", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def cover(doc):
    p(doc, "Capital University", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    p(doc, "Faculty of Computers and Artificial Intelligence", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    p(doc, "Medical Informatics Department", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    for _ in range(3):
        p(doc, "")
    p(doc, "HOLMES IDS", align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True)
    p(doc, "Intrusion Detection System", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    for _ in range(2):
        p(doc, "")
    p(doc, "Supervised by", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    p(doc, "Dr. Soha Ehsan", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    for _ in range(2):
        p(doc, "")
    p(doc, "Implemented by", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    for name in TEAM:
        p(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    for _ in range(2):
        p(doc, "")
    p(doc, "Graduation Project", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    p(doc, "Academic Year 2025-2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    p(doc, "Final Documentation", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
    new_page(doc)
    new_page(doc)


def front_matter(doc):
    h1(doc, "Abstract")
    p(doc, "This graduation project aims to develop HOLMES IDS, a Hybrid Intrusion Detection System (IDS) with a web-based dashboard to enhance network security by monitoring, detecting, explaining, and alerting malicious activities in real time. The system uses signature-based detection to identify known threats and anomaly-based detection to identify unusual patterns in network traffic.")
    p(doc, "The project includes a user-friendly dashboard that provides administrators and analysts with a comprehensive overview of network activity, including real-time alerts, attack summaries, logs, rules, uploaded files, anomaly predictions, explainability results, analytics, and retraining status. The dashboard allows users to filter and review alerts, upload PCAP files for signature analysis, upload CSV feature files for anomaly prediction, manage users and roles, and inspect generated explanations.")
    p(doc, "The IDS is implemented using Python for backend processing, Flask for API services, React and Vite for the frontend dashboard, Scapy for packet capture and packet parsing, SQLite for local data storage, and machine-learning models for anomaly detection. The anomaly engine uses flow-based features, a scaler, an Isolation Forest model, and a stacking classifier. The system is evaluated using project datasets, test PCAP files, CSV files, frontend build tests, backend tests, and manual dashboard workflows.")
    p(doc, "This project contributes to the field of network security by providing an accessible hybrid IDS prototype suitable for academic experimentation and small to medium network-monitoring scenarios. Future work includes improving deployment hardening, expanding real-traffic validation, integrating stronger reporting, and scaling the storage and detection pipeline.")
    new_page(doc)
    new_page(doc)
    h1(doc, "Acknowledgement")
    p(doc, "Our gratitude to those who helped us cannot be put into words. First, we would like to express our appreciation and thankfulness to our supervisor.")
    p(doc, "We would like to express our sincere gratitude to our supervisor, Dr. Soha Ehsan, for her continuous support, valuable guidance, patience, care, and encouragement throughout the duration of this project. Her guidance and mentorship have been critical to our success.")
    p(doc, "We would also like to thank the Faculty of Computers and Artificial Intelligence and the Medical Informatics Department for their academic support during this graduation project.")
    p(doc, "This project would not have been possible without the help, resources, and support of our supervisor, families, colleagues, and everyone who contributed directly or indirectly to this work.")
    new_page(doc)
    new_page(doc)
    h1(doc, "Table of Contents")
    toc = [
        "Chapter 1 Introduction .............................................................................. 12",
        "1.1. Background .................................................................................................. 1",
        "1.1.1. The main area of the project ............................................................. 1",
        "1.1.2. Main techniques and applications .................................................... 1",
        "1.1.3. Motivation and Justification ............................................................. 1",
        "1.2. Problem Statement ....................................................................................... 2",
        "1.3. Objectives .................................................................................................... 2",
        "1.4. Project Scope and Limitations ..................................................................... 3",
        "1.4.1 Project Scope ...................................................................................... 3",
        "1.4.2 Limitations .......................................................................................... 3",
        "1.5. Intrusion Detection System (IDS) overview ............................................... 4",
        "1.5.1. Different Approaches to the IDS ....................................................... 4",
        "1.5.2 Signature-IDS VS Anomaly-IDS ....................................................... 6",
        "1.5.3 What is the best? ............................................................................... 6",
        "1.5.4 Different Forms of the IDS ............................................................... 7",
        "1.6. Project Methodology ................................................................................... 9",
        "1.7. Project Report Outline ............................................................................... 11",
        "Chapter 2 Market and Literature Survey ................................................... 12",
        "2.1 Market Landscape and Challenges ............................................................. 13",
        "2.2 Existing IDS Technologies ......................................................................... 15",
        "2.3 Comparative Evaluation of Techniques ...................................................... 18",
        "Chapter 3 Requirements & Analysis ......................................................... 31",
        "3.1. Project specifications ................................................................................. 20",
        "3.1.1. Stakeholders ..................................................................................... 20",
        "3.1.2. System architecture .......................................................................... 20",
        "3.1.3. Functional Requirements .................................................................. 21",
        "3.1.4. Non-Functional Requirements ........................................................ 23",
        "3.1.5. Use-case Diagram ............................................................................ 25",
        "3.1.6. System Component Diagram ............................................................ 30",
        "3.1.7. Class Diagram .................................................................................. 31",
        "3.1.8. Sequence Diagram ............................................................................ 32",
        "3.1.9. Activity Diagram .............................................................................. 41",
        "Chapter 4 Design, Implementation & Testing .......................................... 45",
        "4.1. System Design ........................................................................................... 46",
        "4.2. Implementation .......................................................................................... 48",
        "4.2.1. Web Dashboard ................................................................................ 49",
        "4.2.2. Database ........................................................................................... 52",
        "4.2.3. Rules ................................................................................................. 55",
        "4.2.4. Deep Packet Inspection .................................................................... 60",
        "4.2.5. Signature-Based IDS ........................................................................ 68",
        "4.2.6. Anomaly-Based IDS ......................................................................... 75",
        "4.2.7. Alerts and Logs .............................................................................. 123",
        "4.2.8. Full Stack Integration & Web Dashboard ...................................... 130",
        "4.3. Testing & Evaluation ................................................................................ 145",
        "4.4. Technologies and Tools ............................................................................ 153",
        "Chapter 5 Results and Discussion ........................................................... 155",
        "Chapter 6 Conclusions and Future Work ................................................ 159",
        "References .............................................................................................. 163",
    ]
    for line in toc:
        p(doc, line, after=1)
    new_page(doc)
    p(doc, "4.2.5.1. Matching Packets with Rules ......................................................... 68")
    p(doc, "4.2.5.2. Signature Rule Matching & Concurrent Packet Processing ......... 73")
    p(doc, "4.2.6.1. Machine Learning Preparations .................................................... 75")
    p(doc, "4.2.6.2. Model Development .................................................................... 101")
    p(doc, "4.2.6.3. Anomaly Detection (Unknown Attacks) ..................................... 111")
    p(doc, "4.2.6.4. Hybrid Detection Method (Stacking + Isolation Forest) ............. 112")
    p(doc, "4.2.6.5. Flow-Based Feature Extraction ................................................... 114")
    p(doc, "4.2.6.6. Anomaly-Based Intrusion Detection System .............................. 119")
    p(doc, "4.3.1. Model Testing ................................................................................ 145")
    p(doc, "4.3.2. DPI Testing .................................................................................... 147")
    p(doc, "4.3.3. Extracting Flow Features ............................................................... 148")
    p(doc, "4.3.4. Performance Evaluation ................................................................ 148")
    p(doc, "4.3.5. Real-Time Detection Testing ........................................................ 149")
    p(doc, "4.3.6. PCAP File Scanning ...................................................................... 150")
    p(doc, "4.3.7. CSV File Prediction ....................................................................... 151")
    p(doc, "4.3.8. Display the Rules to Admin .......................................................... 152")
    p(doc, "4.3.9. Adding Rules through the Dashboard ........................................... 152")
    p(doc, "4.3.10. Limitations ................................................................................... 152")
    new_page(doc)
    new_page(doc)


def chapter_title(doc, chap, title):
    p(doc, f"Chapter {chap}", align=WD_ALIGN_PARAGRAPH.CENTER, size=24, bold=True, before=180, after=20)
    p(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True)
    new_page(doc)


def chapter1(doc):
    chapter_title(doc, 1, "Introduction")
    page_no(doc)
    h2(doc, "1.1. Background")
    h3(doc, "1.1.1. The main area of the project")
    bullet(doc, "The primary focus of this project is the development of an Intrusion Detection System (IDS) designed to monitor, analyze, detect, explain, and alert malicious activities within network traffic.")
    bullet(doc, "Cybersecurity threats have become increasingly sophisticated, making it essential to implement an IDS that can identify potential attacks in real time and alert administrators to take proper actions.")
    bullet(doc, "HOLMES IDS focuses on a hybrid security approach by combining signature-based detection for known threats with anomaly-based detection for unknown or unusual traffic behavior.")
    h3(doc, "1.1.2. Main techniques and applications")
    for item in [
        "React and Vite web dashboard for the frontend.",
        "Flask API for backend processing.",
        "Scapy for packet capture and DPI.",
        "SQLite for storing users, rules, logs, alerts, features, and retraining jobs.",
        "Isolation Forest for unknown attack detection.",
        "Stacking Classifier for known attack classification.",
        "SHAP and LIME for explainability.",
        "Continual learning for model improvement through human labels.",
    ]:
        dot_bullet(doc, item)
    h3(doc, "1.1.3. Motivation and Justification")
    bullet(doc, "Intrusion detection is an essential component of modern security infrastructure for networks and organizations.")
    bullet(doc, "Traditional security measures such as firewalls and antivirus software may not be enough to detect advanced or unknown threats.")
    bullet(doc, "The justification for this project lies in building a practical IDS that can be used to detect known attack signatures, identify anomalous behavior, and present results through an understandable dashboard.")
    page_no(doc)
    h2(doc, "1.2. Problem Statement")
    bullet(doc, "With the rapid expansion of digital networks and online services, cyber threats have become more sophisticated, frequent, and damaging.")
    bullet(doc, "Traditional IDS methods based only on signatures are not always capable of detecting unknown or modified attacks.")
    bullet(doc, "Pure anomaly-based systems can detect suspicious behavior but may produce false positives and require interpretation.")
    bullet(doc, "Therefore, there is a need for a hybrid IDS that combines both techniques and provides clear alerts, logs, dashboard views, analytics, and explanations.")
    h2(doc, "1.3. Objectives")
    for item in [
        "Detect security threats using signature-based rules.",
        "Detect unknown and abnormal network behavior using anomaly-based machine learning.",
        "Provide real-time alerting through live capture.",
        "Support PCAP file scanning and CSV prediction.",
        "Store alerts and logs in a structured database.",
        "Provide dashboards for signature, anomaly, live capture, rules, analytics, and retraining.",
        "Support explainability using SHAP and LIME.",
        "Support role-based access control for different users.",
    ]:
        bullet(doc, item)
    page_no(doc)
    h2(doc, "1.4. Project Scope and Limitations")
    h3(doc, "Project Scope")
    for item in ["Network Traffic Analysis.", "Signature-Based Intrusion Detection.", "Anomaly-Based Intrusion Detection.", "Logging and Alerting System.", "Web-Based Monitoring Interface.", "Explainability and Continual Learning."]:
        bullet(doc, item)
    h3(doc, "Limitations")
    for item in ["Single-node local deployment.", "Live capture depends on Npcap/libpcap availability.", "SQLite may limit scalability under heavy traffic.", "Machine-learning results depend on dataset quality.", "False-positive alarms may still occur.", "Production hardening is required before real deployment."]:
        bullet(doc, item)
    page_no(doc)
    h2(doc, "1.5. Intrusion Detection System (IDS) overview")
    bullet(doc, "IDS operate by continuously monitoring network traffic or system activities and analyzing patterns to detect potential security breaches.")
    bullet(doc, "IDS can provide real-time alerts and actionable insights, allowing administrators to investigate suspicious activity before it escalates.")
    figure_placeholder(doc, "Figure 1.1", None)
    h3(doc, "1.5.1. Different Approaches to the IDS")
    for item in [
        "Signature-Based IDS: compares incoming traffic against known attack signatures.",
        "Anomaly-Based IDS: builds a baseline or uses a model to detect unusual behavior.",
        "Hybrid-Based IDS: combines signature and anomaly methods to provide stronger detection coverage.",
    ]:
        bullet(doc, item)
    page_no(doc)
    h3(doc, "1.5.2 Signature-IDS VS Anomaly-IDS")
    table(doc, ["Key", "Signature IDS", "Anomaly IDS"], [
        ("Advantages", "High accuracy for known threats. Fast and efficient detection.", "Effective against unknown behavior. Adaptable to evolving attacks."),
        ("Limitations", "Weak against zero-day attacks. Requires updated rules.", "Higher false positives. Requires tuning and more resources."),
    ])
    p(doc, "Table 1.1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    h3(doc, "1.5.3 What is the best?")
    dot_bullet(doc, "Signature-based IDS is fast and precise for known attacks.")
    dot_bullet(doc, "Anomaly-based IDS can detect unknown threats but requires more tuning.")
    dot_bullet(doc, "Combining both methods provides stronger protection against cyber threats.")
    h3(doc, "1.5.4 Different Forms of the IDS")
    for item in ["Network-Based IDS monitors network traffic.", "Host-Based IDS monitors activities on one host.", "Hybrid IDS combines multiple monitoring views."]:
        bullet(doc, item)
    page_no(doc)
    h2(doc, "1.6. Project Methodology")
    for item in [
        "Study IDS approaches and project requirements.",
        "Design the system architecture, workflows, and database.",
        "Implement packet parsing, signature detection, anomaly detection, and dashboards.",
        "Integrate uploads, live capture, analytics, explainability, and retraining.",
        "Test the system using PCAP files, CSV files, UI workflows, model evaluation, and build checks.",
        "Document the system design, implementation, results, limitations, and future work.",
    ]:
        bullet(doc, item)
    h2(doc, "1.7. Project Report Outline")
    bullet(doc, "Chapter 2 discusses the market and literature survey.")
    bullet(doc, "Chapter 3 presents requirements and analysis.")
    bullet(doc, "Chapter 4 explains design, implementation, and testing.")
    bullet(doc, "Chapter 5 presents results and discussion.")
    bullet(doc, "Chapter 6 provides conclusions and future work.")


def chapter2(doc):
    chapter_title(doc, 2, "Market and Literature Survey")
    page_no(doc)
    h2(doc, "2.1 Market Landscape and Challenges")
    for text in [
        "The IDS market has grown because organizations need continuous monitoring solutions to protect their assets, services, and sensitive data.",
        "Signature-based IDS tools such as Snort and Suricata remain important because they are effective for known attack patterns.",
        "Anomaly-based techniques are becoming more common because they can identify unusual patterns and potential unknown attacks.",
        "The main challenge is balancing detection accuracy, false positives, performance, explainability, and maintainability.",
    ]:
        bullet(doc, text)
    figure_placeholder(doc, "Figure 2.1")
    page_no(doc)
    h2(doc, "2.2 Existing IDS Technologies")
    for title, lines in [
        ("2.2.1 Snort", ["Snort is a rule-based IDS/IPS widely used for signature-based detection.", "It uses rule headers and rule options to match suspicious traffic."]),
        ("2.2.2 Suricata", ["Suricata is another IDS/IPS engine that supports multi-threaded packet processing.", "It is useful for high-performance network monitoring environments."]),
        ("2.2.3 Zeek", ["Zeek focuses on network security monitoring and event logging.", "It provides rich network analysis rather than only rule matching."]),
        ("2.2.4 Machine Learning IDS", ["Machine-learning IDS systems use features extracted from traffic to classify or detect anomalies.", "They can detect patterns that are difficult to express manually as signatures."]),
    ]:
        h3(doc, title)
        for line in lines:
            bullet(doc, line)
    page_no(doc)
    h2(doc, "2.3 Detection Techniques")
    h3(doc, "2.3.1 Signature-Based IDS")
    bullet(doc, "Signature-based detection is efficient and interpretable but depends on known rules.")
    h3(doc, "2.3.2 Anomaly-Based IDS")
    bullet(doc, "Anomaly detection can identify abnormal behavior but requires careful evaluation to reduce false positives.")
    h3(doc, "2.3.3 Hybrid-Based IDS")
    bullet(doc, "Hybrid IDS combines deterministic signatures with machine-learning anomaly detection.")
    h2(doc, "2.4 Comparative Evaluation of Techniques")
    table(doc, ["Technique", "Strength", "Weakness", "HOLMES IDS Usage"], [
        ("Signature IDS", "Fast and explainable", "Cannot detect all unknown attacks", "Rule matching engine"),
        ("Anomaly IDS", "Detects unusual behavior", "May generate false positives", "Isolation Forest and classifier"),
        ("Hybrid IDS", "Combines both advantages", "More complex implementation", "Main project approach"),
        ("Explainable IDS", "Improves analyst understanding", "Needs stored feature vectors", "SHAP and LIME"),
    ])
    p(doc, "Table 2.1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)


def chapter3(doc):
    chapter_title(doc, 3, "Requirements\n& Analysis")
    page_no(doc)
    h2(doc, "3.1. Project specifications")
    h3(doc, "3.1.1. Stakeholders")
    bullet(doc, "End Users (Network Administrators).")
    bullet(doc, "Cybersecurity Analysts.")
    bullet(doc, "System Integrators and IT Support Teams.")
    bullet(doc, "Security Incident Response Teams.")
    bullet(doc, "Project Supervisor and Academic Evaluators.")
    h3(doc, "3.1.2. System architecture")
    figure_placeholder(doc, "Figure 3.1")
    page_no(doc)
    h3(doc, "3.1.3. Functional Requirements")
    rows = [(f"FR{i:02d}", desc) for i, desc in enumerate([
        "The system should capture and analyze PCAP files using signature-based methods.",
        "The system shall support live capture from a selected network interface.",
        "The system shall support TCP, UDP, ICMP, ARP, and IP packet parsing where available.",
        "The system shall parse and store predefined signature rules.",
        "The system shall apply content, PCRE, flags, dsize, threshold, and detection-filter style checks where implemented.",
        "The administrator shall be able to view and manage rules from the dashboard.",
        "The system shall apply anomaly detection using a trained stacking model and Isolation Forest.",
        "Alerts shall include timestamp, source IP, destination IP, attack type, message, and detection method.",
        "Alerts and logs shall be searchable and viewable from dashboards.",
        "The system shall store all alerts and logs in SQLite.",
        "The system shall provide CSV upload for anomaly prediction.",
        "The system shall provide PCAP upload for signature scanning.",
        "The system shall provide SHAP/LIME explanations for anomaly alerts.",
        "The system shall support user roles and protected routes.",
        "The system shall provide analytics and retraining workflows.",
    ], 1)]
    table(doc, ["Requirement ID", "Description"], rows)
    p(doc, "Table 3.1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    page_no(doc)
    h3(doc, "3.1.4. Non-Functional Requirements")
    table(doc, ["Requirement name", "Description"], [
        ("Usability", "The dashboard shall be intuitive, clear, and easy to use for administrators and analysts."),
        ("Performance & Response Time", "The system should process uploaded files and display results without blocking the dashboard."),
        ("Scalability", "The system should support extending rule sets, adding datasets, and improving models."),
        ("Security", "The interface shall require authentication and role-based access control."),
        ("Maintainability", "The codebase shall be modular, separating rules, packets, logs, alerts, models, APIs, and frontend pages."),
        ("Reliability", "The system must handle malformed PCAP files, corrupted CSV files, and missing packet-capture drivers gracefully."),
    ])
    p(doc, "Table 3.2", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    page_no(doc)
    h3(doc, "3.1.5. Use-case Diagram.")
    h3(doc, "3.1.5.1. Admin")
    p(doc, "This Use Case Diagram illustrates the interaction between the admin and the HOLMES IDS user interface. The system allows the admin to manage users, view alerts, review logs, upload files, manage rules, analyze statistics, and start retraining.")
    table(doc, ["Use Case", "Description"], [
        ("Upload PCAP File", "The admin uploads a .pcap or .pcapng file containing raw network traffic."),
        ("Upload CSV File", "The admin uploads a feature CSV file for anomaly detection."),
        ("View Signature Alerts", "The admin reviews alerts generated by signature-based detection."),
        ("View Anomaly Alerts", "The admin reviews alerts generated by anomaly-based detection."),
        ("Manage Users", "The admin creates, deletes, and updates user roles."),
        ("Start Retraining", "The admin starts a model retraining job after reviewing labeled samples."),
    ])
    p(doc, "Table 3.3", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    figure_placeholder(doc, "Figure 3.2")
    page_no(doc)
    h3(doc, "3.1.5.2. IDS System")
    p(doc, "This Use Case Diagram focuses on internal automated operations handled by the IDS system, including parsing rules, running detection engines, generating logs and alerts, storing features, and analyzing alert statistics.")
    table(doc, ["Use Case", "Description"], [
        ("Analyze Alerts", "The system reviews alerts for frequencies, source IPs, destination IPs, and attack types."),
        ("Parse & Store Rules", "The system extracts rule fields and stores them in the database."),
        ("Start Signature Detection", "The system compares packets with signature rules."),
        ("Start Anomaly Detection", "The system evaluates flow features using the anomaly model."),
        ("Generate Logs", "The system records detection events in the logs table."),
        ("Generate Alerts", "The system records alert events in the alerts table."),
    ])
    p(doc, "Table 3.4", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    figure_placeholder(doc, "Figure 3.3")
    page_no(doc)
    h3(doc, "3.1.6. System Component Diagram")
    figure_placeholder(doc, "Figure 3.4")
    page_no(doc)
    h3(doc, "3.1.7. Class Diagram")
    figure_placeholder(doc, "Figure 3.5")
    page_no(doc)
    h3(doc, "3.1.8. Sequence Diagram")
    for idx, title in enumerate([
        "3.1.8.1. System Initialization",
        "3.1.8.2. Add Rules to Database",
        "3.1.8.3. Detecting Attack with Signature-Based Method",
        "3.1.8.4. Detecting Attack with Anomaly-Based Method",
        "3.1.8.5. CSV Upload Detection",
        "3.1.8.6. PCAP Upload Detection",
        "3.1.8.7. Web UI View",
        "3.1.8.8. Display Rules for Admin",
        "3.1.8.9. Display Alerts for Admin",
        "3.1.8.10. Display Logs for Admin",
    ], 6):
        h3(doc, title)
        figure_placeholder(doc, f"Figure 3.{idx}")
        page_no(doc)
    h3(doc, "3.1.9. Activity Diagram")
    for idx, title in enumerate(["Live Capture Activity", "CSV Prediction Activity", "PCAP Scan Activity", "Retraining Activity"], 16):
        h3(doc, title)
        figure_placeholder(doc, f"Figure 3.{idx}")
        page_no(doc)


def chapter4(doc):
    chapter_title(doc, 4, "Design, Implementation\n& Testing")
    page_no(doc)
    h2(doc, "4.1. System Design")
    bullet(doc, "The system is designed as a full-stack web application with separate frontend, backend, database, detection, explainability, analytics, and retraining components.")
    bullet(doc, "The React frontend communicates with the Flask backend through JSON APIs.")
    bullet(doc, "The backend controls detection logic, uploads, live capture, model inference, database access, and user permissions.")
    figure_placeholder(doc, "Figure 4.1")
    page_no(doc)
    h2(doc, "4.2. Implementation")
    sections = [
        ("4.2.1. Web Dashboard", [
            "The dashboard is implemented using React and Vite. It provides pages for login, signature dashboard, anomaly dashboard, live capture, rules, CSV upload, PCAP upload, analytics, explainability, retraining, and admin management.",
            "Protected routes are used to restrict pages based on user roles.",
        ], "01_signature_dashboard.png"),
        ("4.2.2. Database", [
            "SQLite is used as the local database. It stores rules, logs, alerts, users, alert features, training data, and retrain jobs.",
            "The Database class creates the tables automatically and opens connections with WAL mode.",
        ], None),
        ("4.2.3. Snort / Suricata Rules", [
            "Signature rules are represented with a header and options. The header contains protocol, source, destination, ports, and direction. Options contain message, content, PCRE, flags, size, thresholds, and attack metadata.",
            "The rules are stored in the database and loaded by the signature detection engine.",
        ], "04_rules.png"),
        ("4.2.4. Deep Packet Inspection (DPI)", [
            "Scapy is used to inspect packets and extract protocol, IP addresses, ports, flags, payload, timestamp, and packet length.",
            "The Packet class provides a normalized interface so rule matching can work consistently across packet types.",
        ], None),
        ("4.2.5. Signature-Based IDS", [
            "The signature engine compares each packet with loaded rules. When a rule matches, the system generates a log and alert.",
            "This method is deterministic and suitable for known threats.",
        ], "06_pcap_upload.png"),
        ("4.2.5.1. Matching Packets with Rules", [
            "Packet matching checks protocol, source and destination fields, ports, direction, flags, payload content, regular expressions, and threshold options.",
            "Each matching event includes a message and attack type that can be displayed in the dashboard.",
        ], None),
        ("4.2.5.2. Signature Rule Matching & Concurrent Packet Processing", [
            "PCAP files may contain many packets, so signature analysis should process packets efficiently.",
            "The project uses modular packet parsing and detection functions to keep the implementation maintainable.",
        ], None),
        ("4.2.6. Anomaly-Based IDS", [
            "The anomaly engine uses flow features, a scaler, an Isolation Forest, and a stacking classifier.",
            "The Isolation Forest identifies unknown or abnormal samples, while the classifier predicts known attack labels.",
        ], "02_anomaly_dashboard.png"),
        ("4.2.6.1. Machine Learning Preparations", [
            "The dataset is prepared by cleaning feature columns, aligning the feature order, scaling numeric values, and encoding labels.",
            "The trained artifacts are saved under the Models directory and loaded by the backend during startup.",
        ], None),
        ("4.2.6.2. Model Development", [
            "The stacking classifier combines multiple base learners and a final classifier to improve classification performance.",
            "The project documentation references Decision Tree, Random Forest, Logistic Regression, KNN, CatBoost, and a final meta-classifier.",
        ], None),
        ("4.2.6.3. Anomaly Detection (Unknown Attacks)", [
            "Unknown attack detection depends on the Isolation Forest decision score.",
            "If the score falls below the configured threshold, the sample is treated as anomalous or unknown.",
        ], None),
        ("4.2.6.4. Hybrid Detection Method (Stacking + Isolation Forest)", [
            "The hybrid method combines classifier predictions with outlier detection.",
            "This design improves coverage because it can classify known attacks while still flagging suspicious unknown behavior.",
        ], None),
        ("4.2.6.5. Flow-Based Feature Extraction", [
            "Flow-based features summarize packet behavior over a communication flow.",
            "These features are more suitable for ML prediction than raw packet bytes because they represent behavior patterns.",
        ], None),
        ("4.2.6.6. Anomaly-Based Intrusion Detection System", [
            "The anomaly dashboard shows predictions, alerts, and explanation links.",
            "Feature vectors are stored for SHAP/LIME explanation and for future retraining workflows.",
        ], "05_csv_upload.png"),
        ("4.2.7. If the rule matched with Packet or Detecting Anomalous", [
            "If a signature rule matches or the anomaly engine detects an attack, the system creates alerts and logs in SQLite.",
            "The alert includes timestamp, source IP, destination IP, message, attack label, and detection method.",
        ], None),
        ("4.2.8. AlertAnalyzer Class", [
            "The AlertAnalyzer class calculates simple alert statistics such as top source IP and top destination IP.",
            "These values are used by dashboards to summarize suspicious activity.",
        ], None),
        ("4.2.9. Full Stack Integration & Web Dashboard (Flask Application)", [
            "The Flask backend provides APIs consumed by the React frontend.",
            "The Vite development server proxies API requests to the Flask backend during development.",
        ], "00_login.png"),
    ]
    fig_no = 2
    for title, lines, image in sections:
        page_no(doc)
        h2(doc, title)
        for line in lines:
            bullet(doc, line)
        if image:
            figure_placeholder(doc, f"Figure 4.{fig_no}", image)
        else:
            figure_placeholder(doc, f"Figure 4.{fig_no}")
        fig_no += 1
    page_no(doc)
    h2(doc, "4.3. Testing & Evaluation")
    tests = [
        ("4.3.1. Model Testing", "The model should be evaluated using accuracy, precision, recall, F1-score, confusion matrix, and unknown attack detection."),
        ("4.3.2. DPI Testing", "The Packet class should be tested with TCP, UDP, ICMP, ARP, and malformed packet cases."),
        ("4.3.3. Extracting Flow Features", "Flow feature extraction should be verified against expected feature order."),
        ("4.3.4. Performance Evaluation", "Uploaded files and dashboard operations should complete within acceptable time."),
        ("4.3.5. Real-Time Detection Testing", "Live capture should start correctly when Npcap/libpcap is installed."),
        ("4.3.6. PCAP File Scanning", "PCAP upload should detect signature matches."),
        ("4.3.7. CSV File Prediction", "CSV upload should return anomaly predictions."),
        ("4.3.8. Display the Rules to Admin", "The rules page should display rules loaded from the database."),
        ("4.3.9. Adding Rules through the Dashboard", "Rule management should allow adding and reviewing custom rules where enabled."),
        ("4.3.10. Limitations", "The system limitations must be documented clearly."),
    ]
    for title, body in tests:
        h3(doc, title)
        bullet(doc, body)
        figure_placeholder(doc, f"Figure 4.{fig_no}")
        fig_no += 1
        page_no(doc)
    h2(doc, "4.4. Technologies and Tools")
    for item in [
        "Front-End & UI: React, Vite, Bootstrap, Font Awesome, Chart.js.",
        "Database: SQLite.",
        "DPI: Scapy and Npcap/libpcap.",
        "Programming Language: Python and JavaScript.",
        "Tools: Flask, Flask-Login, pandas, NumPy, scikit-learn, CatBoost, SHAP, LIME, tshark, pytest, npm.",
    ]:
        bullet(doc, item)


def chapter5_6(doc):
    chapter_title(doc, 5, "Results and Discussion")
    page_no(doc)
    h2(doc, "5.1 Findings")
    bullet(doc, "The system successfully integrates signature-based and anomaly-based detection in one web application.")
    bullet(doc, "The dashboards provide practical views for alerts, logs, uploads, live capture, rules, analytics, explainability, and retraining.")
    bullet(doc, "Final measured model values must be inserted after rerunning the evaluation scripts on the final project version.")
    h2(doc, "5.2 Goals Achieved")
    for item in ["Hybrid detection was implemented.", "A web dashboard was implemented.", "SQLite persistence was implemented.", "Live capture support was implemented.", "Explainability and retraining workflows were implemented."]:
        bullet(doc, item)
    h2(doc, "5.3 Further Work")
    for item in ["Improve production security.", "Add more live traffic validation.", "Scale database and storage.", "Add SIEM integration.", "Improve report export.", "Containerize deployment."]:
        bullet(doc, item)
    h2(doc, "5.4 Ethical, Legal, and Social Issues")
    bullet(doc, "Packet capture should only be performed on authorized networks.")
    bullet(doc, "Captured traffic may contain sensitive information and must be handled carefully.")
    bullet(doc, "The project should be used for defensive and educational purposes only.")
    chapter_title(doc, 6, "Conclusions and Future Work")
    page_no(doc)
    h2(doc, "6.1 Conclusions")
    p(doc, "HOLMES IDS demonstrates a hybrid Intrusion Detection System that combines signature-based detection, anomaly-based machine learning, live packet capture, explainability, analytics, role-based access control, and continual learning in one web dashboard.")
    p(doc, "The system provides a practical graduation-project implementation of IDS concepts and shows how different cybersecurity components can be integrated into a single platform.")
    h2(doc, "References")
    refs = [
        "[1] I. Sharafaldin, A. H. Lashkari, and A. A. Ghorbani, CICIDS2017: A Contemporary Dataset for Intrusion Detection, 2018.",
        "[2] M. T. Ribeiro, S. Singh, and C. Guestrin, Why Should I Trust You?: Explaining the Predictions of Any Classifier, 2016.",
        "[3] S. M. Lundberg and S.-I. Lee, A Unified Approach to Interpreting Model Predictions, 2017.",
        "[4] Scapy Documentation.",
        "[5] Flask Documentation.",
        "[6] React Documentation.",
        "[7] scikit-learn Documentation.",
        "[8] Npcap Documentation.",
        "[9] Wireshark/tshark Documentation.",
    ]
    for ref in refs:
        p(doc, ref)


def build():
    doc = Document()
    setup(doc)
    cover(doc)
    front_matter(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5_6(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
