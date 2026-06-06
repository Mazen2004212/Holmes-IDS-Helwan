from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "HOLMES_IDS_Final_Documentation.docx"

COLLEGE_LOGO = Path(r"C:\Users\Dell\Downloads\35835 (300×85) (1).png")
UNIVERSITY_LOGO = Path(r"C:\Users\Dell\Downloads\ؤشحهفشم.png")
SCREENSHOTS = ROOT / "presentation_assets" / "screenshots"


AUTHORS = [
    "Mohamed Ahmed Abdelfattah",
    "Mazen Ibrahim Abdelrazek",
    "Mohamed Abdelgawad Abdelrahman",
    "Hala Mazen Waddad",
    "Sohaila Mustafa Abdelfattah",
]


BLUE = RGBColor(0x00, 0x57, 0x8A)
NAVY = RGBColor(0x0B, 0x25, 0x45)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT = "EAF2F8"


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(rows, headers=None, widths=None):
    doc = table.doc
    total_rows = len(rows) + (1 if headers else 0)
    cols = len(headers or rows[0])
    t = doc.add_table(rows=total_rows, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    start = 0
    if headers:
        for i, h in enumerate(headers):
            set_shading(t.rows[0].cells[i], LIGHT)
            set_cell_text(t.rows[0].cells[i], h, bold=True)
        start = 1
    for r, row in enumerate(rows, start=start):
        for c, val in enumerate(row):
            set_cell_text(t.rows[r].cells[c], val)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    return t


def para(text="", bold=False, italic=False, align=None, size=12):
    doc = para.doc
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    doc = para.doc
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def num(text):
    doc = para.doc
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def heading(text, level=1):
    doc = para.doc
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = NAVY if level == 1 else BLUE
    return p


def chapter_page(title, subtitle=""):
    doc = para.doc
    doc.add_page_break()
    for _ in range(8):
        para("")
    p = para(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=24)
    p.paragraph_format.space_after = Pt(20)
    if subtitle:
        para(subtitle, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
    doc.add_page_break()


def add_figure_placeholder(number, title, filename):
    para(f"[USER-SUPPLIED DIAGRAM HERE - {filename}]", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap = para(f"Figure {number}: {title}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    cap.paragraph_format.space_after = Pt(12)
    para.doc.add_page_break()


def add_image(path, caption, width=6.0):
    doc = para.doc
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        para(caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        doc.add_page_break()


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        styles[name].font.name = "Times New Roman"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].paragraph_format.space_before = Pt(10)
        styles[name].paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("HOLMES IDS - Final Documentation Draft").font.size = Pt(9)


def cover(doc):
    if COLLEGE_LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(COLLEGE_LOGO), width=Inches(3.2))
    if UNIVERSITY_LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(UNIVERSITY_LOGO), width=Inches(1.45))

    para("Capital University", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    para("Faculty of Computers and Artificial Intelligence", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para("Medical Informatics Department", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para("")
    para("HOLMES IDS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=26)
    para("Hybrid Online Learning Model for Enhanced Security", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    para("Intrusion Detection System", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18)
    para("")
    para("Supervised by", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para("Dr. Soha Ehsan", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para("")
    para("Implemented by", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    table([[name] for name in AUTHORS], headers=["Name"])
    para("")
    para("Graduation Project", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para("Academic Year 2025-2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    para("Final Documentation Draft", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_page_break()
    doc.add_page_break()


def front_matter():
    heading("Abstract", 1)
    for text in [
        "This graduation project presents HOLMES IDS, a hybrid Intrusion Detection System designed to monitor, analyze, and detect malicious activities in network traffic. The project combines signature-based detection for known attack patterns with anomaly-based machine learning detection for suspicious or unknown behavior. The system provides a web-based dashboard that helps users review alerts, inspect logs, upload PCAP and CSV files, manage rules, analyze traffic behavior, and evaluate anomaly explanations.",
        "The backend is implemented using Python and Flask, while the frontend is implemented as a React and Vite single-page application. SQLite is used for storing users, rules, logs, alerts, explanation features, training samples, and retraining jobs. Scapy is used for live packet capture and packet parsing, and the anomaly detection workflow uses flow-based features, MinMax scaling, an Isolation Forest, and a trained stacking classifier. The system also supports SHAP and LIME explanations to make anomaly alerts more understandable to analysts.",
        "HOLMES IDS is designed as a practical educational and experimental cybersecurity platform. It demonstrates how signature rules, packet processing, machine learning, explainability, role-based access control, analytics, and continual learning can be integrated in one system. Final measured results must be inserted from the validated project evaluation output before submission."
    ]:
        para(text)
    doc = para.doc
    doc.add_page_break()

    heading("Acknowledgement", 1)
    para("First, we would like to express our sincere gratitude to our supervisor, Dr. Soha Ehsan, for her guidance, support, patience, and valuable feedback throughout the development of this project. Her supervision helped the team refine the project scope, technical direction, and documentation quality.")
    para("We would also like to thank the Faculty of Computers and Artificial Intelligence and the Medical Informatics Department for providing the academic environment that made this graduation project possible.")
    para("Finally, we thank our families, colleagues, and everyone who supported us during the design, implementation, testing, and documentation phases of HOLMES IDS.")
    doc.add_page_break()

    heading("Table of Contents", 1)
    toc = [
        ("Chapter 1 Introduction", "1"),
        ("1.1 Background", "1"),
        ("1.2 Problem Statement", "3"),
        ("1.3 Objectives", "4"),
        ("1.4 Project Scope and Limitations", "5"),
        ("1.5 Intrusion Detection System Overview", "7"),
        ("1.6 Project Methodology", "10"),
        ("Chapter 2 Market and Literature Survey", "12"),
        ("Chapter 3 Requirements & Analysis", "31"),
        ("3.1 Project Specifications", "32"),
        ("3.2 Functional Requirements", "36"),
        ("3.3 Non-Functional Requirements", "38"),
        ("3.4 Use Case Diagrams", "41"),
        ("3.5 System Component Diagram", "43"),
        ("3.6 Class Diagram", "44"),
        ("3.7 Sequence Diagrams", "45"),
        ("3.8 Activity Diagrams", "50"),
        ("Chapter 4 Design, Implementation & Testing", "55"),
        ("4.1 System Design", "56"),
        ("4.2 Implementation", "70"),
        ("4.3 Testing & Evaluation", "110"),
        ("4.4 Technologies and Tools", "130"),
        ("Chapter 5 Results and Discussion", "135"),
        ("Chapter 6 Conclusions and Future Work", "145"),
        ("References", "150"),
        ("Appendices", "155"),
    ]
    table(toc, headers=["Section", "Page"])
    doc.add_page_break()

    heading("List of Figures", 1)
    figs = [
        ("Figure 3.1", "Admin Use Case Diagram"),
        ("Figure 3.2", "IDS System Use Case Diagram"),
        ("Figure 3.3", "System Component Diagram"),
        ("Figure 3.4", "Class Diagram"),
        ("Figure 3.5", "Login Sequence Diagram"),
        ("Figure 3.6", "Live Capture Sequence Diagram"),
        ("Figure 3.7", "CSV Prediction Sequence Diagram"),
        ("Figure 3.8", "PCAP Scan Sequence Diagram"),
        ("Figure 3.9", "Retraining Sequence Diagram"),
        ("Figure 3.10", "Live Capture Activity Diagram"),
        ("Figure 3.11", "ER Diagram"),
        ("Figure 4.1", "High-Level Architecture"),
        ("Figure 4.2", "Detection Pipeline"),
        ("Figure 4.3", "Deployment Diagram"),
    ]
    table(figs, headers=["Figure", "Title"])
    heading("List of Tables", 1)
    tabs = [
        ("Table 1.1", "Signature IDS vs Anomaly IDS"),
        ("Table 2.1", "Comparison of IDS Techniques"),
        ("Table 3.1", "Stakeholders"),
        ("Table 3.2", "Functional Requirements"),
        ("Table 3.3", "Non-Functional Requirements"),
        ("Table 4.1", "Technology Stack"),
        ("Table 4.2", "Database Tables"),
        ("Table 4.3", "Testing Scenarios"),
        ("Table 5.1", "Evaluation Results Template"),
    ]
    table(tabs, headers=["Table", "Title"])
    doc.add_page_break()


def chapter1():
    chapter_page("Chapter 1", "Introduction")
    heading("1.1 Background", 1)
    paragraphs = [
        "Intrusion detection is a critical topic in modern cybersecurity because organizations depend on networks, online services, and interconnected applications to deliver daily operations. As networks become larger and more complex, the attack surface also increases. Attackers can exploit weak services, vulnerable applications, misconfigured devices, and abnormal communication patterns to gain unauthorized access or disrupt services.",
        "The main area of this project is the development of an Intrusion Detection System that monitors network traffic and produces alerts when malicious or suspicious behavior is detected. The project focuses on a hybrid approach because a single detection method is rarely enough. Signature-based detection can efficiently identify known threats using predefined rules, while anomaly-based detection can flag abnormal traffic patterns that may not match a known rule.",
        "HOLMES IDS applies this hybrid concept in a web-based platform. It provides live capture, PCAP scanning, CSV anomaly prediction, dashboards, alert tables, analytics, explainability, and a retraining workflow. The project is implemented as a full-stack system rather than a standalone script, which makes it closer to an operational security tool.",
    ]
    for p in paragraphs:
        para(p)
    heading("1.1.1 The Main Area of the Project", 2)
    for item in [
        "Network traffic monitoring and packet inspection.",
        "Signature-based detection using stored rule definitions.",
        "Flow-based anomaly detection using machine learning.",
        "Web-based dashboards for analysts and administrators.",
        "Storage of logs, alerts, rules, users, model explanation features, and retraining samples.",
    ]:
        bullet(item)
    heading("1.1.2 Main Techniques and Applications", 2)
    for item in [
        "Python and Flask for backend API processing.",
        "React and Vite for the frontend dashboard.",
        "Scapy for packet capture and packet processing.",
        "SQLite for local persistent storage.",
        "Isolation Forest for unknown or anomalous traffic scoring.",
        "Stacking classifier for known attack classification.",
        "SHAP and LIME for explainable anomaly alerts.",
        "Npcap or libpcap for live packet capture support.",
    ]:
        bullet(item)
    heading("1.1.3 Motivation and Justification", 2)
    for p in [
        "The motivation behind HOLMES IDS is the need for a practical and understandable security monitoring tool that can combine rule-based detection with data-driven detection. Traditional tools are useful for known attacks, but they may fail when new traffic behavior does not match a stored signature. Machine learning can help identify suspicious patterns, but its output must be interpreted carefully.",
        "The project is justified by the increasing importance of cybersecurity monitoring in academic, organizational, and medical-informatics contexts. Medical informatics systems handle sensitive data and depend on reliable networked services. A prototype IDS can help demonstrate how traffic monitoring, alerting, explanation, and historical analysis support security awareness and incident response.",
    ]:
        para(p)
    heading("1.2 Problem Statement", 1)
    para("With the expansion of digital networks and web-based services, cyber threats have become more frequent, more varied, and more difficult to detect using a single defensive technique. Firewalls and antivirus tools do not provide complete visibility into packet-level behavior or flow-level anomalies. Signature-based IDS methods detect known attack patterns but may miss unknown attacks, while anomaly-based methods may generate false positives or produce predictions that are difficult to explain.")
    para("The problem addressed by HOLMES IDS is the need to design and implement a hybrid, understandable, and usable IDS platform that can detect known attacks, analyze abnormal traffic, display alerts through a dashboard, and support improvement through human feedback and retraining.")
    heading("1.3 Objectives", 1)
    for item in [
        "Build a web-based IDS dashboard for monitoring alerts, logs, rules, analytics, uploads, and retraining.",
        "Implement signature-based detection for known threats using rule matching.",
        "Implement anomaly-based detection using flow features, Isolation Forest, and a stacking classifier.",
        "Support live packet capture from a selected network interface.",
        "Support offline PCAP scanning and CSV prediction workflows.",
        "Store historical alerts and logs in a structured database.",
        "Provide SHAP and LIME explanations for anomaly alerts where feature vectors are available.",
        "Support role-based access control for administrators and analysts.",
        "Provide an evaluation and testing framework for the project.",
    ]:
        bullet(item)
    heading("1.4 Project Scope and Limitations", 1)
    heading("1.4.1 Project Scope", 2)
    for item in [
        "Hybrid IDS prototype for local deployment.",
        "Signature and anomaly detection workflows.",
        "React frontend and Flask backend integration.",
        "SQLite-based storage for project data.",
        "Live capture where the operating system supports packet capture.",
        "Explainability and continual learning support.",
    ]:
        bullet(item)
    heading("1.4.2 Limitations", 2)
    for item in [
        "SQLite is suitable for the project prototype but not for enterprise-scale event ingestion.",
        "Live capture on Windows requires Npcap with appropriate permissions.",
        "Dataset-based model performance may not equal real live-network performance.",
        "Explainability tools explain classifier behavior and may not fully explain Isolation Forest outlier scores.",
        "Production deployment would require stronger secret management and broader CSRF enforcement.",
    ]:
        bullet(item)
    heading("1.5 Intrusion Detection System Overview", 1)
    para("Intrusion Detection Systems can be categorized by detection approach and monitoring location. The main detection approaches include signature-based IDS, anomaly-based IDS, and hybrid IDS. The main deployment forms include network-based IDS, host-based IDS, and hybrid monitoring solutions.")
    table([
        ("Signature-Based IDS", "Matches known rules", "Fast and interpretable", "Weak against unknown variants"),
        ("Anomaly-Based IDS", "Detects unusual patterns", "Useful for unknown behavior", "Can produce false positives"),
        ("Hybrid IDS", "Combines both methods", "Balanced coverage", "More complex to design and test"),
    ], headers=["Type", "Core Idea", "Advantages", "Limitations"])
    heading("1.6 Project Methodology", 1)
    for item in [
        "Analyze IDS concepts and define project requirements.",
        "Design the system architecture, database schema, and user workflows.",
        "Implement backend modules for authentication, detection, uploads, analytics, explanation, and retraining.",
        "Implement frontend dashboards and protected routes.",
        "Integrate model artifacts and detection pipelines.",
        "Test the system using unit tests, build checks, model evaluation, manual workflows, and uploaded files.",
        "Document results, limitations, and future improvements.",
    ]:
        num(item)
    heading("1.7 Project Report Outline", 1)
    para("Chapter 2 presents the market and literature survey. Chapter 3 presents requirements and analysis. Chapter 4 explains design, implementation, and testing. Chapter 5 discusses results. Chapter 6 provides conclusions and future work.")


def chapter2():
    chapter_page("Chapter 2", "Market and Literature Survey")
    heading("2.1 Market Landscape and Challenges", 1)
    for p in [
        "The IDS market has grown because organizations need tools that can detect intrusions, support compliance, and provide visibility into network behavior. Security teams often use multiple layers of defense, including firewalls, antivirus systems, endpoint detection, SIEM tools, and IDS/IPS solutions.",
        "Signature-based IDS systems remain important because they can detect well-known attacks with low interpretation cost. However, the increasing diversity of attacks creates demand for anomaly detection and behavioral monitoring.",
        "HOLMES IDS is positioned as an academic prototype that demonstrates how these approaches can be combined. It is not designed to compete with enterprise IDS products, but it reflects the architectural principles used in real security-monitoring platforms.",
    ]:
        para(p)
    heading("2.2 Related IDS Concepts", 1)
    for title, body in [
        ("Deep Packet Inspection", "Deep packet inspection examines packet headers and payloads to identify patterns, protocols, and suspicious content. In HOLMES IDS, Scapy is used to parse packets and extract protocol, IP, port, flags, size, and payload information."),
        ("Rule-Based Detection", "Rule-based detection compares packet attributes against predefined rule conditions. This makes the alert logic traceable and understandable."),
        ("Flow-Based Detection", "Flow-based detection groups packets by communication tuple and computes statistical features. These features are used by the anomaly detection model."),
        ("Machine Learning Detection", "Machine learning models can learn complex relationships in feature data. HOLMES IDS uses a stacking classifier and Isolation Forest as saved model artifacts."),
        ("Explainable AI", "Explainable AI helps analysts understand why a model made a prediction. HOLMES IDS uses SHAP and LIME explanations for anomaly alert features."),
    ]:
        heading(f"2.2 - {title}", 2)
        para(body)
    heading("2.3 Comparison of IDS Techniques", 1)
    table([
        ("Snort/Suricata-style signature IDS", "Known attack signatures", "Rule clarity and speed", "Needs updated signatures", "HOLMES IDS uses custom rule objects and stored rules"),
        ("Anomaly-based IDS", "Statistical and ML behavior", "Can detect unusual activity", "Requires careful thresholding", "HOLMES IDS uses Isolation Forest and classifier predictions"),
        ("Hybrid IDS", "Known and unknown detection", "Better coverage", "Higher integration complexity", "HOLMES IDS combines both workflows"),
        ("Explainable IDS", "Human-readable model reasons", "Supports analyst trust", "Can be computationally expensive", "HOLMES IDS stores alert features for SHAP/LIME"),
    ], headers=["Approach", "Focus", "Strength", "Weakness", "Project Relation"])
    heading("2.4 Literature Survey Summary", 1)
    for p in [
        "The literature generally shows that signature detection and anomaly detection solve different parts of the IDS problem. Signature methods are strong when the attack pattern is already known, while anomaly methods are useful when the traffic pattern is unusual but not explicitly covered by a signature.",
        "Hybrid systems are attractive because they reduce dependence on one detection method. However, a hybrid IDS must be carefully designed so that alerts, logs, features, explanations, and retraining data remain consistent.",
        "The literature also emphasizes the need for evaluation. Accuracy alone is not enough; false positives, recall, F1-score, confusion matrix analysis, live-traffic limitations, and interpretability should also be discussed.",
    ]:
        para(p)


def chapter3():
    chapter_page("Chapter 3", "Requirements & Analysis")
    heading("3.1 Project Specifications", 1)
    heading("3.1.1 Stakeholders", 2)
    table([
        ("Admin", "Manage users, view all dashboards, run retraining, access analytics"),
        ("Signature Analyst", "Review signature alerts, PCAP scans, and rules"),
        ("Anomaly Analyst", "Review anomaly alerts, CSV predictions, and explanations"),
        ("Live Operator", "Start and stop live capture and monitor live alerts"),
        ("Supervisor/Evaluator", "Review implementation, design, testing, and documentation"),
    ], headers=["Stakeholder", "Responsibilities"])
    heading("3.1.2 System Architecture", 2)
    para("The system is divided into a React frontend, Flask backend, SQLite database, packet-processing modules, signature detection engine, anomaly detection engine, explainability module, continual learning module, analytics module, uploaded files, and trained model artifacts.")
    add_figure_placeholder("3.3", "System Component Diagram", "figure_3_3_system_component.png")
    heading("3.1.3 Functional Requirements", 2)
    table([
        ("FR-01", "The system shall authenticate users.", "Implemented"),
        ("FR-02", "The system shall enforce role-based access control.", "Implemented"),
        ("FR-03", "The system shall display signature and anomaly dashboards.", "Implemented"),
        ("FR-04", "The system shall allow PCAP upload and signature scanning.", "Implemented"),
        ("FR-05", "The system shall allow CSV upload and anomaly prediction.", "Implemented"),
        ("FR-06", "The system shall support live packet capture.", "Implemented with Npcap/libpcap dependency"),
        ("FR-07", "The system shall store alerts, logs, users, rules, features, and retrain jobs.", "Implemented"),
        ("FR-08", "The system shall provide SHAP/LIME explanations.", "Implemented"),
        ("FR-09", "The system shall support human labels and retraining.", "Implemented"),
        ("FR-10", "The system shall provide analytics queries.", "Implemented"),
    ], headers=["ID", "Requirement", "Status"])
    heading("3.1.4 Non-Functional Requirements", 2)
    table([
        ("Usability", "The dashboard should be clear enough for analysts to review alerts and workflows."),
        ("Maintainability", "Modules should remain separated by responsibility."),
        ("Security", "Passwords should be hashed and routes should enforce roles."),
        ("Performance", "Long operations should not block the dashboard."),
        ("Reproducibility", "Setup, tests, and evaluation should be repeatable."),
        ("Explainability", "Model alerts should include interpretable feature importance when available."),
    ], headers=["Requirement", "Description"])
    heading("3.1.5 Use Case Diagrams", 2)
    add_figure_placeholder("3.1", "Admin Use Case Diagram", "figure_3_1_admin_use_case.png")
    add_figure_placeholder("3.2", "IDS System Use Case Diagram", "figure_3_2_system_use_case.png")
    heading("3.1.6 Class Diagram", 2)
    add_figure_placeholder("3.4", "Class Diagram", "figure_3_4_class_diagram.png")
    heading("3.1.7 Sequence Diagrams", 2)
    for fig, title, file in [
        ("3.5", "Login Sequence Diagram", "figure_3_5_login_sequence.png"),
        ("3.6", "Live Capture Sequence Diagram", "figure_3_6_live_capture_sequence.png"),
        ("3.7", "CSV Prediction Sequence Diagram", "figure_3_7_csv_prediction_sequence.png"),
        ("3.8", "PCAP Scan Sequence Diagram", "figure_3_8_pcap_scan_sequence.png"),
        ("3.9", "Retraining Sequence Diagram", "figure_3_9_retraining_sequence.png"),
    ]:
        add_figure_placeholder(fig, title, file)
    heading("3.1.8 Activity Diagram", 2)
    add_figure_placeholder("3.10", "Live Capture Activity Diagram", "figure_3_10_activity_live_capture.png")
    heading("3.1.9 ER Diagram", 2)
    add_figure_placeholder("3.11", "Database ER Diagram", "figure_3_11_er_diagram.png")


def chapter4():
    chapter_page("Chapter 4", "Design, Implementation & Testing")
    heading("4.1 System Design", 1)
    para("The HOLMES IDS design follows a layered structure. The presentation layer is a React SPA. The application layer is a Flask backend. The detection layer includes packet parsing, signature matching, flow extraction, and anomaly prediction. The persistence layer uses SQLite. The intelligence layer includes explainability, analytics, and continual learning.")
    add_figure_placeholder("4.1", "High-Level Architecture", "figure_4_1_high_level_architecture.png")
    add_figure_placeholder("4.2", "Detection Pipeline", "figure_4_2_detection_pipeline.png")
    add_figure_placeholder("4.3", "Deployment Diagram", "figure_4_3_deployment_diagram.png")
    heading("4.1.1 Frontend Design", 2)
    para("The frontend is organized into pages and shared components. It includes login, signature dashboard, anomaly dashboard, CSV upload, PCAP upload, rules, live capture, admin, explainability, retraining, and analytics pages.")
    screenshots = [
        ("00_login.png", "Figure 4.4: Login Page"),
        ("01_signature_dashboard.png", "Figure 4.5: Signature-Based Dashboard"),
        ("02_anomaly_dashboard.png", "Figure 4.6: Anomaly-Based Dashboard"),
        ("03_live_capture.png", "Figure 4.7: Live Capture Page"),
        ("04_rules.png", "Figure 4.8: Rules Page"),
        ("05_csv_upload.png", "Figure 4.9: CSV Upload Page"),
        ("06_pcap_upload.png", "Figure 4.10: PCAP Upload Page"),
        ("07_analytics.png", "Figure 4.11: Analytics Page"),
        ("08_retrain.png", "Figure 4.12: Retraining Page"),
    ]
    for name, caption in screenshots:
        add_image(SCREENSHOTS / name, caption, width=5.8)
    heading("4.1.2 Backend Design", 2)
    for module, desc in [
        ("UI.py", "Main Flask application entry point, model loading, database initialization, blueprint registration, and legacy routes."),
        ("api_auth.py", "JSON authentication endpoints and CSRF token generation."),
        ("api_routes.py", "Main JSON API endpoints for dashboards, uploads, live capture, admin, explanations, retraining, and analytics."),
        ("auth.py", "User model, password hashing, role definitions, and route decorators."),
        ("DB.py", "SQLite connection handling and table creation."),
        ("app_state.py", "Shared runtime state, paths, model placeholders, rules, and live capture instance."),
    ]:
        heading(f"4.1.2 {module}", 3)
        para(desc)
    heading("4.2 Implementation", 1)
    heading("4.2.1 Packet and Flow Processing", 2)
    for p in [
        "The packet processing workflow uses Scapy packets and wraps them in the Packet class. This class extracts protocol, source IP, destination IP, source port, destination port, TCP flags, payload, packet size, and timestamp.",
        "The flow-processing workflow groups packets into bidirectional flows and computes the features expected by the anomaly model. These features are aligned with the saved feature order before scaling and prediction.",
    ]:
        para(p)
    heading("4.2.2 Database", 2)
    table([
        ("rules", "Stores signature rules."),
        ("logs", "Stores event logs from signature and anomaly detection."),
        ("alerts", "Stores alerts displayed to users."),
        ("users", "Stores username, password hash, and role."),
        ("alert_features", "Stores feature vectors linked to anomaly alerts."),
        ("training_data", "Stores flow features and human labels for retraining."),
        ("retrain_jobs", "Stores retraining status and metrics."),
    ], headers=["Table", "Purpose"])
    heading("4.2.3 Signature-Based IDS", 2)
    for p in [
        "Signature detection compares packet attributes against stored rule definitions. The rule header describes the protocol, source, destination, ports, and direction, while rule options describe payload content, messages, attack names, thresholds, flags, or pattern-matching conditions.",
        "When a packet matches a rule, the system creates a log and alert record. This deterministic path is suitable for detecting known attacks and explaining why the alert was generated.",
    ]:
        para(p)
    heading("4.2.4 Anomaly-Based IDS", 2)
    for p in [
        "The anomaly detection workflow accepts CSV feature files or live/PCAP-derived flow features. Features are scaled using the saved scaler, checked by the Isolation Forest, and classified by the stacking classifier.",
        "The classifier labels known attack categories, while the Isolation Forest helps identify unknown or anomalous behavior. Final evaluation must report validated accuracy, F1, precision, recall, confusion matrix, unknown-detection counts, and full-pipeline behavior.",
    ]:
        para(p)
    heading("4.2.5 Explainability", 2)
    para("The system stores raw feature vectors for anomaly alerts so that explanations can be generated later. SHAP and LIME are used to provide local explanations. These explanations help analysts understand which features influenced the classifier decision.")
    heading("4.2.6 Continual Learning", 2)
    para("The continual learning workflow stores flow features, allows human labels, starts a background retraining job, evaluates the candidate model, and promotes it only if quality gates are passed. This design helps the system adapt while reducing the risk of replacing a model with a weaker candidate.")
    heading("4.3 Testing & Evaluation", 1)
    table([
        ("TC-01", "Login with valid user", "User reaches authorized dashboard", "INSERT"),
        ("TC-02", "Invalid login", "Error message appears", "INSERT"),
        ("TC-03", "PCAP upload", "Signature alerts are returned", "INSERT"),
        ("TC-04", "CSV upload", "Predictions and statistics are returned", "INSERT"),
        ("TC-05", "Live capture without Npcap", "Capture error is shown", "Verified during debugging"),
        ("TC-06", "Live capture with Npcap", "Packet count increases", "INSERT"),
        ("TC-07", "Explain alert", "SHAP/LIME output appears", "INSERT"),
        ("TC-08", "Retraining", "Job status and metrics update", "INSERT"),
        ("TC-09", "Frontend build", "Vite build succeeds", "INSERT"),
        ("TC-10", "Backend tests", "pytest passes", "INSERT"),
    ], headers=["ID", "Scenario", "Expected Result", "Actual Result"])
    heading("4.4 Technologies and Tools", 1)
    table([
        ("Frontend", "React, Vite, React Router, Bootstrap, Font Awesome, Chart.js"),
        ("Backend", "Python, Flask, Flask-Login, Werkzeug"),
        ("Database", "SQLite"),
        ("Packet Processing", "Scapy, Npcap/libpcap"),
        ("Machine Learning", "scikit-learn, CatBoost, pandas, NumPy, joblib"),
        ("Explainability", "SHAP, LIME, matplotlib"),
        ("TLS Analysis", "tshark/Wireshark where installed"),
        ("Testing", "pytest, Vite build, manual workflows"),
    ], headers=["Category", "Tools"])


def chapter5_6():
    chapter_page("Chapter 5", "Results and Discussion")
    heading("5.1 Findings", 1)
    para("The final findings must be filled after the team reruns the evaluation scripts and stores the results. Existing repository documentation reports high performance on CIC-style evaluation data, but the final thesis should not present those values as final unless they are reproduced.")
    table([
        ("Stacking classifier accuracy", "About 99.74%", "INSERT validated result"),
        ("Weighted F1", "About 99.74%", "INSERT validated result"),
        ("Weighted precision", "About 99.74%", "INSERT validated result"),
        ("Weighted recall", "About 99.74%", "INSERT validated result"),
        ("Full pipeline attack detection rate", "About 99.85%", "INSERT validated result"),
        ("Full pipeline benign accuracy", "About 94.35%", "INSERT validated result"),
        ("Unknown attack detection", "100/100 simulated samples", "INSERT validated result"),
    ], headers=["Metric", "Repository-Reported Value", "Final Validated Value"])
    heading("5.2 Goals Achieved", 1)
    for item in [
        "The system integrates signature-based and anomaly-based detection.",
        "The frontend provides dashboards and workflows for major project features.",
        "The backend exposes APIs for authentication, uploads, live capture, analytics, explanations, and retraining.",
        "The database stores operational and learning data.",
        "The system supports explainable anomaly alerts and continual learning.",
    ]:
        bullet(item)
    heading("5.3 Discussion", 1)
    para("Strong benchmark results would indicate that the saved model artifacts perform well on the evaluation data. However, real live traffic can differ from CIC-style datasets. The discussion must therefore separate validated dataset performance from expected deployment behavior.")
    heading("5.4 Ethical, Legal, and Social Issues", 1)
    para("Packet capture must be performed only on networks where the team has authorization. Captured traffic can contain sensitive information, so storage, screenshots, logs, and reports must be handled carefully. The project should be used for defensive education, monitoring, and research, not unauthorized surveillance.")
    chapter_page("Chapter 6", "Conclusions and Future Work")
    heading("6.1 Conclusions", 1)
    para("HOLMES IDS demonstrates the design and implementation of a hybrid intrusion detection system that combines signature matching, anomaly detection, live capture, explainability, analytics, role-based access, and continual learning in one web-based platform.")
    para("The project provides a practical environment for understanding IDS concepts and evaluating how different detection methods can cooperate. The main contribution is the integration of multiple cybersecurity and machine-learning workflows into a usable graduation-project system.")
    heading("6.2 Future Work", 1)
    for item in [
        "Deploy the system with production-grade secret management.",
        "Apply CSRF protection consistently to all mutating API routes.",
        "Add more benchmark datasets and live-network validation scenarios.",
        "Improve model drift monitoring and retraining governance.",
        "Move from SQLite to a scalable database for larger deployments.",
        "Add report export, alert severity scoring, and SIEM integration.",
        "Containerize backend and frontend deployment.",
    ]:
        bullet(item)
    heading("References", 1)
    refs = [
        "[1] I. Sharafaldin, A. H. Lashkari, and A. A. Ghorbani, \"CICIDS2017: A Contemporary Dataset for Intrusion Detection,\" 2018. VERIFY final bibliographic details.",
        "[2] M. T. Ribeiro, S. Singh, and C. Guestrin, \"Why Should I Trust You?: Explaining the Predictions of Any Classifier,\" 2016.",
        "[3] S. M. Lundberg and S.-I. Lee, \"A Unified Approach to Interpreting Model Predictions,\" NeurIPS, 2017.",
        "[4] Scapy official documentation, packet manipulation and sniffing reference.",
        "[5] Flask official documentation.",
        "[6] React official documentation.",
        "[7] scikit-learn official documentation for IsolationForest and ensemble models.",
        "[8] Npcap official documentation for Windows packet capture support.",
        "[9] Wireshark/tshark official documentation.",
    ]
    for r in refs:
        para(r)


def appendices():
    chapter_page("Appendices", "Supporting Project Material")
    heading("Appendix A - Repository Manifest", 1)
    files = [
        ("UI.py", "Flask application entry point"),
        ("api_auth.py", "Authentication API"),
        ("api_routes.py", "Main JSON API routes"),
        ("DB.py", "Database schema and connection wrapper"),
        ("auth.py", "Users and role-based access control"),
        ("packet.py", "Packet parsing"),
        ("flow.py", "Flow feature extraction"),
        ("rule.py", "Signature rule object"),
        ("signature_IDS.py", "Signature-based IDS"),
        ("anomaly_IDS.py", "Anomaly-based IDS"),
        ("live_capture.py", "Live capture and background processing"),
        ("explainability.py", "SHAP/LIME explanations"),
        ("continual_learning.py", "Retraining workflow"),
        ("analytics.py", "Analytics query builder"),
        ("frontend/src", "React frontend application"),
        ("Models", "Saved models, scaler, encoders, datasets"),
        ("UML", "Existing diagram source files"),
    ]
    table(files, headers=["File/Folder", "Purpose"])
    heading("Appendix B - Required System Design Images", 1)
    para("The following images should be generated from SYSTEM_DESIGN_MERMAID.md and inserted into the final document.")
    required = [
        "figure_3_1_admin_use_case.png",
        "figure_3_2_system_use_case.png",
        "figure_3_3_system_component.png",
        "figure_3_4_class_diagram.png",
        "figure_3_5_login_sequence.png",
        "figure_3_6_live_capture_sequence.png",
        "figure_3_7_csv_prediction_sequence.png",
        "figure_3_8_pcap_scan_sequence.png",
        "figure_3_9_retraining_sequence.png",
        "figure_3_10_activity_live_capture.png",
        "figure_3_11_er_diagram.png",
        "figure_4_1_high_level_architecture.png",
        "figure_4_2_detection_pipeline.png",
        "figure_4_3_deployment_diagram.png",
    ]
    for r in required:
        bullet(r)
    heading("Appendix C - Installation Guide", 1)
    for step in [
        "Install Python 3.10 or compatible version.",
        "Create and activate a virtual environment.",
        "Install backend requirements using pip install -r requirements.txt.",
        "Install Node.js and npm.",
        "Run npm install inside the frontend folder.",
        "Install Npcap on Windows with WinPcap API-compatible mode for live capture.",
        "Start the backend using python UI.py.",
        "Start the frontend using npm run dev.",
        "Open the frontend URL and login using the configured admin account.",
    ]:
        num(step)
    heading("Appendix D - Notes for Final Submission", 1)
    for item in [
        "Replace placeholders marked INSERT or VERIFY.",
        "Insert rendered system design diagrams.",
        "Insert final validated testing screenshots and command outputs.",
        "Verify references in IEEE format.",
        "Confirm university cover-page requirements.",
        "Change default admin credentials before any public demonstration.",
    ]:
        bullet(item)
    heading("Appendix E - Detailed Module Documentation", 1)
    module_details = [
        ("UI.py", [
            "The UI.py module is the main runtime entry point of the backend application. It creates the Flask application object, configures login management, loads machine-learning artifacts, registers API blueprints, initializes database tables, loads rules, and starts the development server.",
            "The file also contains legacy server-rendered routes in addition to the JSON API used by the React frontend. This transitional structure is useful for development but should be refactored in future work so that API logic, page rendering, and startup code remain separated.",
            "During startup, UI.py loads the stacking classifier, Isolation Forest, scaler, label encoder, and feature order. These objects are stored in app_state so they can be used by upload routes, live capture, anomaly prediction, explainability, and retraining workflows."
        ]),
        ("api_auth.py", [
            "The api_auth.py module implements JSON authentication routes under /api/auth. It provides the CSRF token route, login route, logout route, and current-user status route.",
            "The login route receives username and password data from the React frontend, checks the password against the stored hash, creates a Flask-Login session, and returns the authenticated user role and default frontend route.",
            "This module is important because every dashboard workflow depends on user identity and role. Future hardening should apply CSRF checks consistently across all mutating backend endpoints."
        ]),
        ("api_routes.py", [
            "The api_routes.py module is the central API layer for the React single-page application. It exposes dashboard data, upload workflows, live capture control, admin user management, explanation requests, retraining operations, and analytics queries.",
            "This module converts backend objects such as Alert, Log, Rule, and User into JSON dictionaries suitable for frontend display. It also coordinates file validation and calls the relevant detection modules.",
            "The live capture endpoints maintain a reference to the current LiveCapture instance through app_state. Recent error handling stores and displays capture failures such as missing Npcap/WinPcap support."
        ]),
        ("DB.py", [
            "The DB.py module wraps SQLite connection creation and table initialization. It creates tables for rules, logs, alerts, users, alert features, training data, and retraining jobs.",
            "SQLite is configured with WAL mode and a timeout. This is appropriate for a local prototype, although a production IDS would require a more scalable storage architecture.",
            "The clear_table method uses an allowlist of table names to prevent arbitrary table deletion, which is an important defensive programming practice."
        ]),
        ("auth.py", [
            "The auth.py module defines user roles and access control behavior. The roles include admin, signature_analyst, anomaly_analyst, and live_operator.",
            "Passwords are hashed using Werkzeug security utilities. The role_required decorator protects backend routes by checking whether the current user's role is allowed to access a route.",
            "A default admin account is created when no users exist. This supports first-run usability, but the final report should clearly state that default credentials must be changed before any real deployment."
        ]),
        ("packet.py", [
            "The packet.py module converts Scapy packet objects into a normalized Packet abstraction. This abstraction extracts protocol, source and destination IP addresses, ports, flags, payload, timestamp, and packet size.",
            "Normalizing packets is necessary because signature rules should not depend directly on Scapy's internal representation. The Packet class gives the rule engine a predictable interface.",
            "The module also handles different protocols and malformed payloads gracefully, which is important when analyzing real network traffic or PCAP files."
        ]),
        ("rule.py and RuleProcessor.py", [
            "The rule.py module represents one signature rule and implements matching logic. It supports header fields and rule options such as content, PCRE, flags, size, thresholds, and attack messages.",
            "RuleProcessor.py is responsible for loading and seeding default rules. The rule database allows rules to be displayed, reused, and managed through the system.",
            "This rule-processing layer is the foundation of the signature-based IDS workflow."
        ]),
        ("signature_IDS.py", [
            "The signature_IDS.py module applies signature rules to packet data. For PCAP analysis, it can process many packets and return alerts for matched rules.",
            "Signature-based alerts are deterministic because each alert can be traced back to a rule condition. This makes signature detection easier to explain than many machine-learning predictions.",
            "The module is used by both offline PCAP upload workflows and live capture processing."
        ]),
        ("flow.py", [
            "The flow.py module groups packets into network flows and computes the feature vector expected by the anomaly model.",
            "Flow-based features summarize communication behavior across packets rather than inspecting each packet independently. This is useful for detecting attacks whose behavior emerges over multiple packets.",
            "The feature order must match the trained model artifacts exactly; otherwise, predictions can become invalid."
        ]),
        ("anomaly_IDS.py", [
            "The anomaly_IDS.py module performs anomaly prediction over uploaded CSV feature files and flow features. It aligns feature columns, scales the data, evaluates Isolation Forest scores, and applies the stacking classifier.",
            "The Isolation Forest flags unknown or anomalous patterns, while the classifier predicts known labels. The combination supports both known-class detection and unknown-attack handling.",
            "Prediction results should be evaluated with accuracy, precision, recall, F1-score, confusion matrix, and unknown-detection analysis."
        ]),
        ("live_capture.py", [
            "The live_capture.py module captures packets in the background, applies signature detection, buffers packets for flow processing, writes training features, and maintains live status counters.",
            "The design uses multiple threads: a capture loop, a flow-processing loop, and a batch writer loop. This prevents all work from being performed in the packet callback.",
            "Live capture depends on operating-system packet capture support. On Windows, Npcap must be installed, preferably with WinPcap API compatibility enabled."
        ]),
        ("explainability.py", [
            "The explainability.py module generates SHAP and LIME explanations for anomaly alerts where feature vectors were stored.",
            "Explainability improves analyst trust by showing which features influenced a prediction. However, explanations must be interpreted carefully because they may explain the classifier more directly than the Isolation Forest outlier score.",
            "The final documentation should include screenshots of explanation outputs after the team generates real anomaly alerts."
        ]),
        ("continual_learning.py", [
            "The continual_learning.py module implements the human-in-the-loop retraining workflow. It loads stored training samples, applies human labels, retrains candidate models, evaluates metrics, and promotes models only when quality gates are satisfied.",
            "This workflow is valuable because network behavior changes over time. However, retraining must be governed carefully so that incorrect labels do not degrade the model.",
            "The final report should include a tested retraining run or clearly mark retraining results as pending."
        ]),
        ("analytics.py", [
            "The analytics.py module implements a query-builder style analytics engine. It supports filtering, grouping, metrics, time ranges, and pattern-oriented analysis without exposing raw SQL to the user.",
            "This design is safer than allowing arbitrary SQL queries from the frontend. It also gives analysts a structured way to inspect alerts, logs, rules, users, and training data."
        ]),
        ("Frontend React Application", [
            "The frontend is organized into pages, components, API utilities, and authentication context. Protected routes restrict access based on the authenticated user role.",
            "The interface provides separate pages for signature detection, anomaly detection, live capture, CSV upload, PCAP upload, rules, admin, explainability, retraining, and analytics.",
            "Screenshots from the presentation assets are included in Chapter 4 and should be replaced with final high-resolution screenshots before submission if needed."
        ]),
    ]
    for module, paragraphs in module_details:
        heading(module, 2)
        for text in paragraphs:
            para(text)
        para.doc.add_page_break()

    heading("Appendix F - Extended Testing Checklist", 1)
    tests = [
        ("Authentication", "Verify valid login, invalid login, logout, session persistence, and unauthorized route handling."),
        ("Role-Based Access", "Verify admin, signature analyst, anomaly analyst, and live operator access boundaries."),
        ("Signature Dashboard", "Verify alert table loading, top source IP, top destination IP, rule count, and empty-state behavior."),
        ("Anomaly Dashboard", "Verify anomaly alerts, stored feature vectors, and explanation links."),
        ("CSV Upload", "Verify valid CSV, invalid extension, missing file, large file, unknown attack labeling, and prediction truncation."),
        ("PCAP Upload", "Verify valid PCAP, invalid extension, empty file, signature matches, and optional TLS metadata."),
        ("Live Capture", "Verify no-driver error, valid driver start, stop, packet count, alert count, and interface display."),
        ("Rules", "Verify default rules load after fresh database initialization and are visible in the dashboard."),
        ("Explainability", "Verify SHAP and LIME outputs for stored anomaly alerts."),
        ("Retraining", "Verify sample listing, label update, retrain start, progress status, metrics, promotion, and rollback behavior."),
        ("Analytics", "Verify filters, grouping, time ranges, and pattern queries."),
        ("Frontend Build", "Verify npm run build succeeds without errors."),
        ("Backend Tests", "Verify pytest tests pass in the correct virtual environment."),
    ]
    table(tests, headers=["Area", "Checklist"])


def build():
    global doc
    doc = Document()
    para.doc = doc
    table.doc = doc
    setup_doc(doc)
    cover(doc)
    front_matter()
    chapter1()
    chapter2()
    chapter3()
    chapter4()
    chapter5_6()
    appendices()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
